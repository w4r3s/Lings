import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
import pandas as pd
from docx import Document
from docx.shared import RGBColor
import logging
from plugins.deepl_translator import translate_text  # 假设deepl_translator插件存在

# 配置日志记录
logging.basicConfig(filename='translation.log', level=logging.DEBUG, format='%(asctime)s:%(levelname)s:%(message)s')

class TextEditor:
    def __init__(self, root, untranslated_segments=None, file_path=None):
        self.root = root
        self.root.title("Text Editor")

        self.file_path = tk.StringVar(value=file_path)
        self.data = pd.DataFrame(columns=["Original", "Translation"])
        self.font_size = tk.IntVar(value=10)
        self.editing_item = None
        self.editing_column = None
        self.entry_edit = None

        self.create_widgets()
        if file_path:
            self.load_word_file()
        elif untranslated_segments:
            self.load_untranslated_segments(untranslated_segments)

    def create_widgets(self):
        frame = tk.Frame(self.root)
        frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

        tk.Label(frame, text="Word File:").grid(row=0, column=0, sticky="w")
        tk.Entry(frame, textvariable=self.file_path, width=50).grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        tk.Button(frame, text="Load", command=self.load_file).grid(row=0, column=2, padx=5, pady=5)

        tree_frame = tk.Frame(frame)
        tree_frame.grid(row=1, column=0, columnspan=3, padx=5, pady=5, sticky="nsew")

        self.tree = ttk.Treeview(tree_frame, columns=("Original", "Translation"), show="headings", selectmode="browse")
        self.tree.heading("Original", text="Original")
        self.tree.heading("Translation", text="Translation")
        self.tree.column("Original", anchor="w", width=300)
        self.tree.column("Translation", anchor="w", width=300)

        self.vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscroll=self.vsb.set, xscroll=self.hsb.set)

        self.vsb.pack(side="right", fill="y")
        self.hsb.pack(side="bottom", fill="x")
        self.tree.pack(fill="both", expand=True)

        style = ttk.Style()
        style.configure("Treeview.Heading", font=("Arial", 10, "bold"))
        self.update_treeview_style()

        self.tree.bind("<Double-1>", self.on_double_click)
        self.tree.bind("<Return>", self.on_return_key)
        self.tree.bind("<Escape>", self.on_escape_key)
        self.tree.bind("<Button-3>", self.show_context_menu)  # 绑定右键点击事件
        self.tree.bind("<Motion>", self.on_motion)  # 绑定鼠标移动事件
        self.vsb.bind("<Motion>", self.on_motion)  # 绑定垂直滚动条的鼠标移动事件
        self.tree.bind("<Configure>", self.on_resize)  # 绑定窗口大小调整事件

        frame.grid_rowconfigure(1, weight=1)
        frame.grid_columnconfigure(1, weight=1)

        button_frame = tk.Frame(self.root)
        button_frame.pack(padx=10, pady=5, fill=tk.X)

        tk.Button(button_frame, text="Add Entry", command=self.add_entry).grid(row=0, column=0, padx=5, pady=5)
        tk.Button(button_frame, text="Delete Entry", command=self.delete_entry).grid(row=0, column=1, padx=5, pady=5)
        tk.Button(button_frame, text="Save", command=self.save_file).grid(row=0, column=2, padx=5, pady=5)
        tk.Button(button_frame, text="Save As", command=self.save_file_as).grid(row=0, column=3, padx=5, pady=5)
        tk.Button(button_frame, text="Translate Selected with DeepL", command=self.translate_selected).grid(row=0, column=4, padx=5, pady=5)

        tk.Label(button_frame, text="Font Size:").grid(row=1, column=0, padx=5, pady=5)
        tk.Spinbox(button_frame, from_=8, to=20, textvariable=self.font_size, command=self.update_treeview_style).grid(row=1, column=1, padx=5, pady=5)

        # 创建右键菜单
        self.context_menu = tk.Menu(self.tree, tearoff=0)
        self.context_menu.add_command(label="Translate Selected (DeepL)", command=self.translate_selected)
        # 这里可以添加更多插件的选项
        self.context_menu.add_separator()
        self.context_menu.add_command(label="Add Entry", command=self.add_entry)
        self.context_menu.add_command(label="Delete Entry", command=self.delete_entry)

    def show_context_menu(self, event):
        try:
            item = self.tree.identify_row(event.y)
            if item:
                self.tree.selection_set(item)
                self.context_menu.post(event.x_root, event.y_root)
        finally:
            self.context_menu.grab_release()

    def update_treeview_style(self):
        font_size = self.font_size.get()
        style = ttk.Style()
        style.configure("Treeview", font=("Arial", font_size), rowheight=font_size + 10)
        style.configure("Treeview.Heading", font=("Arial", font_size, "bold"))
        style.map("Treeview", background=[('selected', 'gray90')], foreground=[('selected', 'black')])
        self.tree.tag_configure('oddrow', background='white')
        self.tree.tag_configure('evenrow', background='lightblue')
        for col in self.tree["columns"]:
            self.tree.heading(col, text=col, anchor="w")

    def load_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if file_path:
            self.file_path.set(file_path)
            self.load_word_file()
            logging.info(f'Loaded file: {file_path}')

    def load_word_file(self):
        file_path = self.file_path.get()
        if not file_path:
            return

        logging.debug(f'Loading Word file: {file_path}')
        doc = Document(file_path)
        rows = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                logging.debug(f'Paragraph text: {text}')
                original = text
                translation = ''
                rows.append({"Original": original, "Translation": translation})

        self.data = pd.DataFrame(rows, columns=["Original", "Translation"])
        if self.data.empty:
            logging.warning("No entries found in the document.")
        self.update_treeview()

    def load_untranslated_segments(self, segments):
        rows = [{"Original": segment, "Translation": ""} for segment in segments]
        self.data = pd.DataFrame(rows, columns=["Original", "Translation"])
        if self.data.empty:
            logging.warning("No entries found in the document.")
        self.update_treeview()

    def update_treeview(self):
        for row in self.tree.get_children():
            self.tree.delete(row)
        for i, (_, row) in enumerate(self.data.iterrows()):
            tags = ('evenrow',) if i % 2 == 0 else ('oddrow',)
            self.tree.insert("", "end", iid=i, values=(row["Original"], row.get("Translation", "")), tags=tags)

    def add_entry(self):
        original = simpledialog.askstring("Add Entry", "Enter the original text:")
        translation = simpledialog.askstring("Add Entry", "Enter the translation:")
        if original and translation:
            self.data = pd.concat([self.data, pd.DataFrame([{"Original": original, "Translation": translation}])], ignore_index=True)
            self.update_treeview()
            logging.info(f'Added entry: {original} -> {translation}')

    def delete_entry(self):
        selected_item = self.tree.selection()
        if selected_item:
            original = self.tree.item(selected_item, "values")[0]
            self.data = self.data[self.data["Original"] != original]
            self.update_treeview()
            logging.info(f'Deleted entry: {original}')

    def save_file(self):
        file_path = self.file_path.get()
        if file_path:
            self.prompt_and_save(file_path)
        else:
            messagebox.showerror("Save File", "No file path specified.")
            logging.error('No file path specified for saving translated document.')

    def save_file_as(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if file_path:
            self.prompt_and_save(file_path)
            self.file_path.set(file_path)

    def prompt_and_save(self, file_path):
        save_original = messagebox.askyesno("Save Original", "Do you want to save the original text as well?")
        self.save_word_file(file_path, save_original)
        messagebox.showinfo("Save File", "Translated document saved successfully.")
        logging.info(f'Saved translated document: {file_path}')

    def save_word_file(self, file_path, save_original):
        doc = Document()
        for _, row in self.data.iterrows():
            original = row["Original"].strip()
            translation = row["Translation"].strip()
            if original:
                p = doc.add_paragraph()
                if save_original:
                    p.add_run(original).font.color.rgb = RGBColor(0, 0, 255)  # 原文设置为蓝色
                    if translation:
                        p.add_run("\n" + translation).font.color.rgb = RGBColor(255, 0, 0)  # 翻译设置为红色
                else:
                    if translation:
                        p.add_run(translation).font.color.rgb = RGBColor(255, 0, 0)  # 翻译设置为红色
        doc.save(file_path)

    def translate_selected(self):
        selected_item = self.tree.selection()
        if selected_item:
            original_text = self.tree.item(selected_item, "values")[0]
            translated_text = translate_text(original_text, "EN")  # 假设目标语言是英文，你可以根据需要调整
            if translated_text:
                self.tree.set(selected_item, column=1, value=translated_text)
                original = self.tree.item(selected_item, "values")[0]
                self.data.loc[self.data["Original"] == original, "Translation"] = translated_text
                logging.info(f'Translated text: {original_text} -> {translated_text}')
            else:
                messagebox.showerror("Translation Error", "Failed to translate the selected text.")

    def on_double_click(self, event):
        if not self.tree.selection():
            return
        self.editing_item = self.tree.selection()[0]
        self.editing_column = self.tree.identify_column(event.x)
        self.start_edit()

    def on_return_key(self, event):
        self.save_edit()

    def on_escape_key(self, event):
        self.cancel_edit()

    def on_motion(self, event):
        if self.entry_edit:
            self.update_edit_box_position()

    def on_resize(self, event):
        if self.entry_edit:
            self.update_edit_box_position()

    def start_edit(self):
        if not self.editing_column:
            return
        item = self.tree.item(self.editing_item, "values")
        column = int(self.editing_column[1:]) - 1  # column name to index
        if column < 0:
            return
        bbox = self.tree.bbox(self.editing_item, self.editing_column)
        if not bbox:
            return
        x, y, width, height = bbox
        self.entry_edit = tk.Entry(self.tree, width=width // 10)
        self.entry_edit.insert(0, item[column])
        self.entry_edit.select_range(0, tk.END)
        self.entry_edit.place(x=x, y=y, width=width, height=height)
        self.entry_edit.focus_set()

    def save_edit(self):
        new_value = self.entry_edit.get()
        column = int(self.editing_column[1:]) - 1
        self.tree.set(self.editing_item, column=column, value=new_value)
        original = self.tree.item(self.editing_item, "values")[0]
        self.data.loc[self.data["Original"] == original, self.tree["columns"][column]] = new_value
        logging.info(f'Edited entry: {original} -> {new_value}')
        self.entry_edit.destroy()
        self.entry_edit = None
        self.update_treeview()

    def cancel_edit(self):
        if self.entry_edit:
            self.entry_edit.destroy()
            self.entry_edit = None

    def update_edit_box_position(self):
        if not self.editing_item or not self.editing_column:
            return
        bbox = self.tree.bbox(self.editing_item, self.editing_column)
        if not bbox:
            return
        x, y, width, height = bbox
        self.entry_edit.place(x=x, y=y, width=width, height=height)

def open_editor(untranslated_segments=None, file_path=None):
    editor_root = tk.Toplevel()
    TextEditor(editor_root, untranslated_segments, file_path)
    editor_root.mainloop()

if __name__ == "__main__":
    root = tk.Tk()
    app = TextEditor(root)
    root.mainloop()
