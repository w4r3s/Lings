import json
import os
import logging
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

# 配置日志记录
logging.basicConfig(filename='translation.log', level=logging.INFO, format='%(asctime)s:%(levelname)s:%(message)s')

class Translator:
    def __init__(self, language, category_path, strict_punctuation=True, ignore_case=False, partial_match=False):
        self.language = language
        self.category_path = category_path
        self.strict_punctuation = strict_punctuation
        self.ignore_case = ignore_case
        self.partial_match = partial_match
        self.is_arabic = 'ar' in category_path
        self.word_dict = self.load_word_dict()

    def load_word_dict(self):
        try:
            with open(self.category_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                logging.info(f'Loaded word dictionary from {self.category_path}')
                logging.debug(f'Word dictionary content: {data}')
                return data.get("translations", {})
        except FileNotFoundError:
            logging.error(f'Word dictionary file {self.category_path} not found')
            return {}

    def save_word_dict(self):
        with open(self.category_path, 'w', encoding='utf-8') as f:
            json.dump({"translations": self.word_dict}, f, ensure_ascii=False, indent=4)
        logging.info(f'Saved word dictionary to {self.category_path}')

    def add_translation(self, word, translation):
        self.word_dict[word] = translation
        self.save_word_dict()
        logging.info(f'Added translation: {word} -> {translation}')

    def translate_word(self, word):
        if self.is_arabic:
            word_clean = ''.join(filter(str.isalnum, word)) if not self.strict_punctuation else word

            if self.partial_match:
                for key in self.word_dict:
                    if key in word_clean:
                        translated_word = word.replace(key, self.word_dict[key])
                        logging.info(f'Partially matched and translated: {word} -> {translated_word}')
                        return translated_word

            translation = self.word_dict.get(word_clean, word)
            if translation != word:
                logging.info(f'Translated: {word} -> {translation}')
            return translation
        else:
            lookup_word = word.lower() if self.ignore_case else word

            if self.partial_match:
                for key in self.word_dict:
                    key_compare = key.lower() if self.ignore_case else key
                    if key_compare in lookup_word:
                        translated_word = lookup_word.replace(key_compare, self.word_dict[key])
                        logging.info(f'Partially matched and translated: {lookup_word} -> {translated_word}')
                        return translated_word

            if self.strict_punctuation:
                translation = self.word_dict.get(lookup_word, word)
            else:
                word_clean = ''.join(filter(str.isalnum, lookup_word))
                translation = self.word_dict.get(word_clean, lookup_word)
            
            if lookup_word != translation:
                logging.info(f'Translated: {word} -> {translation}')
            
            return translation

    def translate_text(self, text):
        logging.debug(f'Translating text: {text}')
        if text in self.word_dict:
            translated_text = self.word_dict[text]
            logging.info(f'Translated paragraph: {text} -> {translated_text}')
            return translated_text, True
        
        words = text.split()
        new_words = [self.translate_word(word) for word in words]
        translated_text = ' '.join(new_words)
        
        if translated_text != text:
            logging.info(f'Translated paragraph: {text} -> {translated_text}')
            return translated_text, True
        else:
            return text, False

    def translate_paragraph(self, para):
        new_text, modified = self.translate_text(para.text)
        if modified:
            para.clear()  # 清除原始段落内容
            run = para.add_run(new_text)  # 添加新内容
            if self.is_arabic:
                para.paragraph_format.alignment = 3  # 从右到左对齐
                run.rtl = True  # 设置从右到左
                run.font.name = 'Arial'  # 设置阿拉伯语字体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            if not self.strict_punctuation:
                run.font.color.rgb = RGBColor(255, 0, 0)  # 设置颜色为红色

    def translate_document(self, input_file, output_file, progress_callback=None):
        doc = Document(input_file)
        logging.info(f'Loaded document: {input_file}')
        
        total_elements = len(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                total_elements += len(row.cells)

        processed_elements = 0
        untranslated_segments = []

        def update_progress():
            nonlocal processed_elements
            processed_elements += 1
            if progress_callback:
                progress_callback(processed_elements / total_elements * 100)

        for para in doc.paragraphs:
            new_text, modified = self.translate_text(para.text)
            if modified:
                para.clear()  # 清除原始段落内容
                para.add_run(new_text)  # 添加新内容
                if not self.strict_punctuation:
                    para.runs[-1].font.color.rgb = RGBColor(255, 0, 0)  # 设置颜色为红色
            else:
                untranslated_segments.append(para.text)
            update_progress()
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        new_text, modified = self.translate_text(para.text)
                        if modified:
                            para.clear()  # 清除原始段落内容
                            para.add_run(new_text)  # 添加新内容
                            if not self.strict_punctuation:
                                para.runs[-1].font.color.rgb = RGBColor(255, 0, 0)  # 设置颜色为红色
                        else:
                            untranslated_segments.append(para.text)
                        update_progress()

        doc.save(output_file)
        logging.info(f'Saved translated document: {output_file}')

        return untranslated_segments
